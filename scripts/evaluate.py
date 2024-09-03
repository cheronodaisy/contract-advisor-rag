import os
from dotenv import load_dotenv
from langchain.prompts import ChatPromptTemplate
from langchain.schema.runnable import RunnablePassthrough
from langchain.schema.output_parser import StrOutputParser
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.retrievers.multi_query import MultiQueryRetriever
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from matplotlib.colors import LinearSegmentedColormap
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain.retrievers import EnsembleRetriever
from langchain_community.retrievers import BM25Retriever


load_dotenv()

llm = ChatOpenAI(
        model_name="gpt-4o-mini",
        api_key=os.getenv("OPENAI_API_KEY"),
        temperature=0,
        max_tokens=800,
        model_kwargs={"top_p": 0, "frequency_penalty": 0, "presence_penalty": 0},
    )
#text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
embedding = OpenAIEmbeddings()

def extract_text_from_docx(file_path):
    text = ""
    doc = DocxDocument(file_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text
def extract_text_from_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

file_paths = ["/home/daisy/Desktop/tenx/ContractAdvisorRAG/samples/Sample_Contract.docx"]
texts = []

for file_path in file_paths:
    if file_path.endswith('.pdf'):
        texts.append(extract_text_from_pdf(file_path))
    elif file_path.endswith('.docx'):
        texts.append(extract_text_from_docx(file_path))


chunk_size = 500
chunk_overlap = 50
text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

chunks = []
for text in texts:
    chunks.extend(text_splitter.split_text(text))

documents = [Document(page_content=chunk) for chunk in chunks]

# bm25_retriever = BM25Retriever.from_documents(documents)
# bm25_retriever.k = 2

# chroma_vectorstore = Chroma.from_documents(documents, embedding)
# chroma_retriever = chroma_vectorstore.as_retriever()

# ensemble_retriever = EnsembleRetriever(
#     retrievers=[bm25_retriever, chroma_retriever], weights=[0.5, 0.5]
# )

# multiquery
vectorstore = Chroma.from_documents(documents=documents, embedding=embedding)
retriever = MultiQueryRetriever.from_llm(
    retriever=vectorstore.as_retriever(), llm=llm)

rag_chain = RetrievalQA.from_chain_type(llm, chain_type="stuff", retriever=retriever)

#RAGAS
from datasets import Dataset

questions = [
    "What should be included in the invoice?", 
    "What happens if a contractor breaches the contract?",
    "What is the deadline for submitting the invoicing schedule?",
    "When will Mercy Corps submit payment?",
    "Who is responsible for paying taxes in the agreement?"
]
ground_truths = [
    "Each invoice will include (i) the Contract Number; (ii) Contractor’s name and address; (iii) a description of the Services performed, (iv) the dates such Services were performed, (v) a pricing calculation based on the payment terms, (vi) properly reimbursable expenses (if any) incurred along with receipts for such expenses (if applicable) for all individual expenses exceeding $25 USD, and (vii) such other information as Mercy Corps may reasonably request.",
    "In the event of termination due to Contractor’s breach of this Contract or by Contractor for Contractor’s convenience, Mercy Corps will not be obligated to pay Contractor for any partially completed work. ",
    "Final invoices must be submitted within 60 days of the end date of the Contract. Contractor recognizes that in many cases Mercy Corps’ donor will not reimburse Mercy Corps for invoices submitted beyond 60 days after the termination of a contract and therefore Mercy Corps will have no obligation to pay any portion of invoices received more than 60 days after the end date of the Contract.",
    "Except as otherwise provided in the Statement of Services, Mercy Corps will pay each invoice (or adjusted invoice if the subject of dispute) in accordance with the Payment Terms within 30 days after the later of (i) receipt of the invoice or (ii) resolution of the items set forth in the notice of disputed charges.",
    "Except as otherwise provided in the Statement of Services, Contractor is responsible for all expenses incurred by it in performing under this Contract and all taxes, duties and other governmental charges with respect to the provision of Services.  If the law requires Mercy Corps to withhold taxes from payments to Contractor, Mercy Corps may withhold those taxes and pay them to the appropriate taxing authority.  Mercy Corps will deliver to Contractor an official notice for such taxes.  Mercy Corps will use reasonable efforts to minimize any taxes withheld to the extent allowed by law."
]

answers = []
contexts = []

# Inference
for query in questions:
    result = rag_chain.invoke(query)
    answer = result['result'] if 'result' in result else result  # Ensure result is in the expected format
    relevant_docs = retriever.invoke(query)
    answers.append(answer)
    contexts.append([doc.page_content for doc in relevant_docs])

# To dict
data = {
    "question": questions,
    "answer": answers,
    "contexts": contexts,
    "ground_truth": ground_truths
}
# Convert dict to dataset
dataset = Dataset.from_dict(data)

# Print dataset structure for debugging
#print(dataset[2])

from ragas import evaluate
from ragas.metrics import (
    faithfulness,
    answer_relevancy,
    context_recall,
    context_precision,
)

result = evaluate(
    dataset = dataset, 
    metrics=[
        context_precision,
        context_recall,
        faithfulness,
        answer_relevancy,
    ],
)

df = result.to_pandas()
print(df.head())
mean_df=df[['faithfulness','answer_relevancy', 'context_precision', 'context_recall']].mean(axis=0)
mean_df.to_csv('mean_metric_ensemble.csv, encoding="utf-8", index=True')
heatmap_data = df[['context_precision', 'context_recall', 'faithfulness', 'answer_relevancy']]

cmap = LinearSegmentedColormap.from_list('green_red', ['red', 'green'])

plt.figure(figsize=(10, 8))
sns.heatmap(heatmap_data, annot=True, fmt=".2f", linewidths=.5, cmap=cmap)

plt.yticks(ticks=range(len(df['question'])), labels=df['question'], rotation=0)

plt.show()
