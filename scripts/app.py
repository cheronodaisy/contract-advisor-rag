import streamlit as st
import os
from dotenv import load_dotenv
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain.retrievers import EnsembleRetriever
from langchain_community.retrievers import BM25Retriever
import fitz
from langchain.docstore.document import Document
import io
from docx import Document as DocxDocument
import clipboard

load_dotenv()

llm = ChatOpenAI(
    model_name="gpt-4o-mini",
    api_key=os.getenv("OPENAI_API_KEY"),
    temperature=0,
    model_kwargs={"top_p": 0, "frequency_penalty": 0, "presence_penalty": 0},
)

embedding = OpenAIEmbeddings()

def extract_text_from_pdf(file):
    text = ""
    try:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
    return text

def extract_text_from_docx(file):
    text = ""
    try:
        doc = DocxDocument(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
    return text

def extract_text_from_txt(file):
    text = ""
    try:
        text = file.read().decode("utf-8")
    except Exception as e:
        st.error(f"Error reading TXT file: {e}")
    return text

def process_documents(texts):
    chunk_size = 500
    chunk_overlap = 50

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    chunks = []
    for text in texts:
        chunks.extend(text_splitter.split_text(text))

    documents = [Document(page_content=chunk) for chunk in chunks]

    bm25_retriever = BM25Retriever.from_documents(documents)
    bm25_retriever.k = 2

    chroma_vectorstore = Chroma.from_documents(documents, embedding)
    chroma_retriever = chroma_vectorstore.as_retriever()

    ensemble_retriever = EnsembleRetriever(
        retrievers=[bm25_retriever, chroma_retriever], weights=[0.5, 0.5]
    )

    rag_chain = RetrievalQA.from_chain_type(llm, chain_type="stuff", retriever=ensemble_retriever)
    return rag_chain

# Initialize session state
if "rag_chain" not in st.session_state:
    st.session_state.rag_chain = None
if "history" not in st.session_state:
    st.session_state.history = []
if "copied" not in st.session_state:
    st.session_state.copied = []

def rag_qa(query):
    if st.session_state.rag_chain:
        response = st.session_state.rag_chain.invoke(query)
        return response['result']
    return "No contract loaded."

def on_copy_click(text):
    st.session_state.copied.append(text)
    clipboard.copy(text)

# Sidebar for file upload
with st.sidebar:
    st.header("Upload Contract")
    uploaded_files = st.file_uploader("Choose files", type=["pdf", "docx", "txt"], accept_multiple_files=True)
    
    if uploaded_files:
        texts = []
        for file in uploaded_files:
            if file.type == "application/pdf":
                file_text = extract_text_from_pdf(file)
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_text = extract_text_from_docx(file)
            elif file.type == "text/plain":
                file_text = extract_text_from_txt(file)
            else:
                st.error(f"Unsupported file type: {file.type}")
                continue
            
            if file_text:
                texts.append(file_text)
        if texts:
            st.session_state.rag_chain = process_documents(texts)
            st.write("Contracts processed and ready for Q&A.")
        else:
            st.error("No text extracted from uploaded files.")

# Main chatbox
st.title("Your Contract Assistant")

query = st.text_input("Enter your question")

if st.button("Send"):
    if query:
        answer = rag_qa(query)
        st.session_state.history.insert(0, ("Assistant", answer))
        st.session_state.history.insert(0, ("You", query))
        st.write("## Chat")
        for sender, message in st.session_state.history:
            if sender == "You":
                st.markdown(f"<div style='text-align: right; padding: 10px; background-color: #AEBACF; border: 1px solid #a1d8a4; color: #333; border-radius: 10px; margin-bottom: 5px;'>{message}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='text-align: left; padding: 10px; background-color: #f3f4f6; color: #333; border-radius: 10px; margin-bottom: 5px; border: 1px solid #e0e0e0;'>{message}</div>", unsafe_allow_html=True)
        st.button("📋", on_click=on_copy_click, args=(answer,))
        st.toast("Copied to clipboard", icon='✅')

st.write("Disclaimer! Please verify any important information as the system may make mistakes.")
